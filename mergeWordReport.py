#!/usr/bin/python3
# _*_ coding: utf-8 _*_
import datetime
import docx
import time
from docxOperate import merge2TeamReport,replaceText

def getTraveleNum(filePath):
    doc = docx.Document(filePath)
    table= doc.tables[0]
    count=0
    for row in table.rows:
        travelCell=row.cells[2]
        if not ('无' in travelCell.text):
            count=count+1
    return count-1

def getLeaveNum(filePath):
    doc = docx.Document(filePath)
    table= doc.tables[0]
    count=0
    for row in table.rows:
        travelCell=row.cells[1]
        if not ('无' in travelCell.text):
            count=count+1
    return count-1

def createTeamReport(resultDir,orgName):
    templateName='template/report_template.docx'
    #读取模板文件
    templateFile = docx.Document(templateName)   
    monthStr=str(datetime.datetime.now().month)
    timeStr= time.strftime("%Y%m%d", time.localtime())   
    teamReportName=resultDir+'/'+orgName+timeStr+"_汇总.docx"
    templateFile.save(teamReportName)
    # 将周报模板复制到临时目录下
    return teamReportName

def mergeWordReport(resultDir,totalNum,orgName,reportsNameList):
    teamReportName=createTeamReport(resultDir,orgName)
    doneDict={}
    for reportName in reportsNameList:
        merge2TeamReport(reportName,teamReportName)
    replaceText("totalNum",totalNum,teamReportName)
    leaveNum=getLeaveNum(teamReportName)
    replaceText("leaveNum",leaveNum,teamReportName)
    traveleNum=getTraveleNum(teamReportName)
    replaceText("travelNum",traveleNum,teamReportName)
    return teamReportName

