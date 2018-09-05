import docx
def replaceText(textName,textValue,filePath):
    doc = docx.Document(filePath)
    markStr='${'+textName+'}$'
    for paragraph in doc.paragraphs:
        if markStr in paragraph.text:
            oText=paragraph.text
            paragraph.text =oText.replace(markStr,str(textValue))
    doc.save(filePath)
    return
def merge2TeamReport(fileName,teamReportName):
    doc = docx.Document(fileName)
    personalTables = doc.tables    
    mergeDoc = docx.Document(teamReportName)
    teamWorkTables = mergeDoc.tables
    # 合并出勤信息
    merge2TeamTables(personalTables, teamWorkTables,0)
    # 合并工作完成信息
    merge2TeamTables(personalTables, teamWorkTables,2)
     # 合并待协调问题
    merge2AllTeamTables(personalTables, teamWorkTables,3)
    mergeDoc.save(teamReportName)
def merge2TeamTables(personalTables, teamWorkTables,tableNum):
    personalWorkTable = personalTables[tableNum]
    selectedCells = personalWorkTable.rows[1].cells
    i = 0
    cells = teamWorkTables[tableNum].add_row().cells
    for cell in selectedCells:
        text=cell.text
        array=text.split('\n')
        j=0
        combinedText=''
        for words in array:
            flag1='无'.endswith(words)
            flag2=len(words)==0
            flag3='-' in words
            if (not flag1) and (not flag2) and (not flag3):
                j=j+1
                combinedText=combinedText+str(j)+'.'+words+'\n'
            else:
                combinedText=combinedText+words+'\n'
        cells[i].text = combinedText
        i = i + 1
def merge2AllTeamTables(personalTables, teamWorkTables,tableNum):
    personalWorkTable = personalTables[tableNum]
    for row in personalWorkTable.rows:
        selectedCells = row.cells
        i = 0
        cells = teamWorkTables[tableNum].add_row().cells
        for cell in selectedCells:
            text=cell.text
            array=text.split('\n')
            j=0
            combinedText=''
            for words in array:
                flag1='无'.endswith(words)
                flag2=len(words)==0
                flag3='-' in words
                if (not flag1) and (not flag2) and (not flag3):
                    j=j+1
                    combinedText=combinedText+str(j)+'.'+words+'\n'
                else:
                    combinedText=combinedText+words+'\n'
            cells[i].text = combinedText
            i = i + 1
