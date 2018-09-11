#!/usr/bin/python3
# _*_ coding: utf-8 _*_
import docx
import time
from xlsOperate import merge2HistoryXlsx, mergeAllinfo2HistoryXlsx
from fileUtil import scanDir,getModifyTime


def extractDoneInfo(filePath):
    doc = docx.Document(filePath)
    # 获取工作表格
    table = doc.tables[2]
    rows = table.rows
    count = 0
    doneDict = {}
    for row in rows:
        if count > 0:
            name = row.cells[0].text
            # 获取第三栏完成情况
            done = row.cells[2].text
            doneDict.update({name: done})
        count = count+1
    return doneDict

def extractAllInfo(filePath):
    modifyTime=getModifyTime(filePath)
    timeStr=time.strftime("%Y%m%d",time.localtime(modifyTime)) 
    doc = docx.Document(filePath)
    # 获取工作表格
    table = doc.tables[2]
    rows = table.rows
    count = 0
    allDict = {}
    for row in rows:
        if count>0:
            name = row.cells[0].text
            # 获取第三栏完成情况
            compareInfo = row.cells[1].text
            doneInfo = row.cells[2].text
            planInfo = row.cells[3].text
            allDict.update({name: [ timeStr,compareInfo, doneInfo, planInfo]})
        count=count+1
    return allDict


def mergeExcelReport(resultDir, orgName, histroryFileName):
    fileList = scanDir(resultDir)
    for fileName in fileList:
        if fileName.find(".docx") > 0:
            # 将从原始文件获取周报信息改为从合并后的周报获取信息
            doneDict = extractDoneInfo(fileName)
            histroryFileName = merge2HistoryXlsx(
                resultDir, orgName, doneDict, histroryFileName)
    return histroryFileName


def mergeOrg2DepartmentReport(resultDir, reportDirName, orgName, histroryFileName):
    fileList = scanDir(reportDirName)
    for fileName in fileList:
        if fileName.find(".docx") > 0:
            # 将从班组周报信息改为从合并后的周报获取信息
            singleWeekDict = extractAllInfo(fileName)
            histroryFileName= mergeAllinfo2HistoryXlsx(resultDir, orgName, singleWeekDict, histroryFileName)

